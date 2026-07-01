import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

print("Generating Formal Structured Investment Memo...")

doc = docx.Document()

# Base Formatting Setup
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# ==========================================
# 1. COVER PAGE
# ==========================================
p_cover_top = doc.add_paragraph()
p_cover_top.paragraph_format.space_before = Pt(100)

title = doc.add_heading('INSTITUTIONAL INVESTMENT MEMO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = subtitle.add_run('FUNDAMENTAL VALUATION & STOCHASTIC ANALYSIS')
run_sub.font.size = Pt(14)
run_sub.font.bold = True
run_sub.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph().paragraph_format.space_after = Pt(120)

# Metadata Table on Cover Page
table_meta = doc.add_table(rows=5, cols=2)
table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Target Entity:", "Sun Pharmaceutical Industries Ltd. (NSE: SUN.NS)"),
    ("Sector / Industry:", "Indian Pharmaceuticals (Nifty Pharma Index)"),
    ("Lead Valuation Analyst:", "Bijesh Singha"),
    ("Valuation Date:", "July 2026"),
    ("Document Classification:", "CONFIDENTIAL - INSTITUTIONAL RESEARCH ONLY")
]

for i, (label, val) in enumerate(meta_data):
    row = table_meta.rows[i]
    r0 = row.cells[0].paragraphs[0].add_run(label)
    r0.bold = True
    r0.font.size = Pt(11)
    r1 = row.cells[1].paragraphs[0].add_run(val)
    r1.font.size = Pt(11)

doc.add_page_break()

# ==========================================
# 2. EXECUTIVE SUMMARY PAGE
# ==========================================
doc.add_heading('1. EXECUTIVE SUMMARY', level=1)

p_recom = doc.add_paragraph()
p_recom.paragraph_format.space_before = Pt(6)
p_recom.add_run('INVESTMENT RECOMMENDATION: ').bold = True
run_sell = p_recom.add_run('SELL / UNDERWEIGHT')
run_sell.bold = True
run_sell.font.size = Pt(13)
run_sell.font.color.rgb = RGBColor(220, 38, 38) # Bold Red

doc.add_paragraph('Based on our deterministic Driver-Based Discounted Cash Flow (DCF) model and 10,000-iteration Monte Carlo stochastic simulation, Sun Pharmaceutical Industries (SUN.NS) is currently trading at a severe premium relative to its underlying cash-generative capacity.')

# Valuation Summary Table
table_exec = doc.add_table(rows=6, cols=2)
table_exec.alignment = WD_TABLE_ALIGNMENT.CENTER
exec_data = [
    ("Valuation Metric", "Value (INR / Share)"),
    ("Current Market Trading Price", "₹1,868.00"),
    ("Base Case DCF Intrinsic Value (50th Percentile)", "₹887.87"),
    ("Implied Target Downside", "-52.47%"),
    ("Bear Case Stress Value (10th Percentile)", "₹780.22"),
    ("Bull Case Stress Value (90th Percentile)", "₹1,006.34")
]

for i, (col1, col2) in enumerate(exec_data):
    row = table_exec.rows[i]
    c0 = row.cells[0]
    c1 = row.cells[1]
    c0.paragraphs[0].text = col1
    c1.paragraphs[0].text = col2
    if i == 0:
        set_cell_background(c0, "1E293B")
        set_cell_background(c1, "1E293B")
        c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        c0.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        c1.paragraphs[0].runs[0].font.bold = True
    elif i == 2: # Highlight Base Case
        c0.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 38, 38)

doc.add_heading('Core Investment Thesis', level=2)
p_thesis = doc.add_paragraph()
p_thesis.add_run('While Sun Pharma exhibits stable top-line expansion (~10% CAGR) and robust EBIT operating margins (~23%), the required capital reinvestment rate—specifically the Working Capital drag and maintenance Capital Expenditures required to sustain pharmaceutical supply chains—absorbs approximately 30% of operating profits. ')
p_thesis.add_run('The live equity market is pricing Sun Pharma for absolute operational perfection, completely ignoring the balance sheet reinvestment burden required to support future revenue growth.')

doc.add_page_break()

# ==========================================
# 3. ASSUMPTIONS (LIST DOWN EVERYTHING)
# ==========================================
doc.add_heading('2. COMPLETE VALUATION ASSUMPTIONS & PARAMETERS', level=1)
doc.add_paragraph('To ensure full mathematical transparency and institutional auditability, every parameter, proxy, and macroeconomic assumption utilized across our deterministic and stochastic modeling engines is detailed below.')

doc.add_heading('2.1 Macroeconomic & Cost of Capital (WACC) Assumptions', level=2)
table_wacc = doc.add_table(rows=7, cols=3)
wacc_assumptions = [
    ("Parameter", "Assumed Value", "Financial Rationale & Source"),
    ("Risk-Free Rate (Rf)", "7.00%", "Benchmark India 10-Year Government Bond Yield."),
    ("Equity Risk Premium (ERP)", "6.00%", "Consensus long-term equity risk premium for Indian equity markets."),
    ("Equity Beta (β)", "0.75", "LSEG live extraction or defensive proxy beta for large-cap pharmaceuticals."),
    ("Cost of Debt (Pre-Tax Rd)", "8.00%", "Annualized quarterly interest expense divided by total debt outstanding."),
    ("Effective Corporate Tax Rate (Tc)", "25.00%", "Normalized statutory Indian corporate tax rate proxy."),
    ("Calculated WACC", "11.45%", "Dynamic weighted blend based on live E/V and D/V capital structure.")
]
for i, row_data in enumerate(wacc_assumptions):
    row = table_wacc.rows[i]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.paragraphs[0].text = val
        if i == 0:
            set_cell_background(cell, "334155")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_heading('2.2 Base Case Deterministic 5-Year Forecast Drivers', level=2)
table_drivers = doc.add_table(rows=5, cols=3)
driver_assumptions = [
    ("Forecast Driver", "Baseline Parameter", "Operational Justification"),
    ("Baseline Revenue (Year 0)", "₹450,000 Million", "Normalized trailing historical base revenue established in Phase 2."),
    ("Top-Line Revenue Growth Rate", "10.00% p.a.", "Historical industry growth ceiling for large-cap mature pharma manufacturers."),
    ("Operating Profit Margin (EBIT)", "23.00%", "Verified operating margin efficiency derived from historical financial statements."),
    ("Reinvestment Rate (CapEx + NWC)", "30.00% of EBIT", "Capital required to maintain physical plants and fund inventory/receivables.")
]
for i, row_data in enumerate(driver_assumptions):
    row = table_drivers.rows[i]
    for j, val in enumerate(row_data):
        cell = row.cells[j]
        cell.paragraphs[0].text = val
        if i == 0:
            set_cell_background(cell, "334155")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_heading('2.3 Terminal Valuation Assumptions', level=2)
doc.add_paragraph('• Perpetual Gordon Growth Rate (g): 5.00% (Reflects long-term Indian nominal GDP growth ceiling).')
doc.add_paragraph('• Terminal Year Reinvestment: Assumed to normalize in perpetuity as reinvestment matches depreciation.')

doc.add_heading('2.4 Stochastic Monte Carlo Simulation Parameters', level=2)
doc.add_paragraph('• Simulation Iterations: 10,000 independent randomized trials.')
doc.add_paragraph('• Revenue Growth Distribution: Normal Distribution | Mean = 10.0% | Standard Deviation = 2.0%.')
doc.add_paragraph('• EBIT Margin Distribution: Normal Distribution | Mean = 23.0% | Standard Deviation = 1.5%.')
doc.add_paragraph('• Net Working Capital Drag Mechanics: Programmatically coupled to top-line revenue expansion to penalize artificial cash inflation.')

doc.add_page_break()

# ==========================================
# 4. QUANTITATIVE FINDINGS & REVERSE DCF
# ==========================================
doc.add_heading('3. MONTE CARLO STOCHASTIC FINDINGS', level=1)
doc.add_paragraph('Stochastic stress testing confirms that even across 10,000 iterations simulating potential market upside and operating outperformance, Sun Pharma fails to reach current market trading levels:')
doc.add_paragraph('• 90th Percentile (Bull Case): ₹1,006.34 per share')
doc.add_paragraph('• 50th Percentile (Base Case): ₹887.87 per share')
doc.add_paragraph('• 10th Percentile (Bear Case): ₹780.22 per share')

doc.add_heading('4. REVERSE DCF: THE MARKET DELUSION', level=1)
doc.add_paragraph('To mathematically justify the current trading price of ₹1,868 under our verified WACC (11.45%) and margin framework (23%), Sun Pharma would have to achieve a Compound Annual Growth Rate (CAGR) of 37.3% every single year for the next 5 years.')
doc.add_paragraph('Because Sun Pharma is a ₹1.8+ Trillion market-cap incumbent whose historical organic growth ceiling averages ~10-12%, the market\'s implied expectation of 37.3% continuous growth is physically and structurally impossible to execute.')

doc.add_paragraph('_' * 70)
p_footer = doc.add_paragraph('CONFIDENTIAL - FOR INSTITUTIONAL DISTRIBUTION ONLY')
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('Investment_Memo_SUN_NS.docx')
print("Successfully generated Investment_Memo_SUN_NS.docx with Cover Page, Executive Summary, and Detailed Assumptions.")

